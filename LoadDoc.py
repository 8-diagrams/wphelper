
from pydocx import PyDocX
import docx 

def toHtml(docf):
    print("proc ", docf)
    f = open(docf, "rb")
    data = f.read()
    f.close()
    print("data-len " , len(data) )
    # Pass in a path
    html = PyDocX.to_html(docf)
    return html 

def toHtml2(docf):
    import docx2html 
    html = docx2html.convert(docf)
    return html
    
from spire.doc import *
from spire.doc.common import *

def conv3(docf, outf):


    # Create a Document instance
    document = Document()

    # Load a Word document
    document.LoadFromFile(docf)

    # Set the type of CSS style sheet as internal
    document.HtmlExportOptions.CssStyleSheetType = CssStyleSheetType.Internal

    # Embed images in HTLM code
    document.HtmlExportOptions.ImageEmbedded = True

    # Export form fields as plain text
    document.HtmlExportOptions.IsTextInputFormFieldAsText = True

    # Save the document as an HTML file
    document.SaveToFile(outf, FileFormat.Html)
    document.Close()

def toHtml4(file_path):
    doc = docx.Document(file_path)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    return text

import poword

def doc2docx11(input_path, output_path):
    poword.doc2docx(input_path, output_path)

import textract
def toHtml8(docf):
    text = textract.process(docf)
    return text 

import win32com
import win32com.client

def doc2docx( doc , docx ):
    wps=win32com.client.Dispatch('Word.application')
    doc1=wps.Documents.Open(doc, ReadOnly=True)
    import time 
    #time.sleep(4)
    doc1.SaveAs(FileName=docx,FileFormat=12)

import subprocess

def excuteCommand(com):
    
    ex = subprocess.Popen(com, stdout=subprocess.PIPE)
    out, err  = ex.communicate()
    status = ex.wait()
    #print("cmd in:", com)
    import guesscs 
    code = guesscs.cs(out)
    stext = out.decode( encoding=code)
    #print(f"cmd out with code {code} : ", stext )
    return stext

def getTextFromDoc(docf):
    import subprocess
    cmd_fullpath = 'C:\\env\\ahome\\antiword\\antiword.exe'
    cmd = f'{cmd_fullpath} -m UTF-8.txt -t "{docf}" '
    print("CMD->", cmd)
    try:
        out_bytes = excuteCommand(cmd)
    except Exception as e:
        import traceback
        #print( traceback.format_exc() )
        #print( "exception", e )
        return ""
    #out_text = out_bytes.decode('utf-8')
    out_bytes = out_bytes.replace("\a","")
    return out_bytes

def getTextFromDoc_Libre(docpath):
    import subprocess 
    cmd_fullpath = 'C:\\env\\LibreOffice\\program\\soffice --cat '
    cmd = f'{cmd_fullpath} "{docpath}" '
    print("CMD->", cmd)
    try:
        out_bytes = excuteCommand(cmd)
    except Exception as e:
        import traceback
        print( "Excep:", traceback.format_exc() )
        print( "exception as", e )
        return ""
    #out_text = out_bytes.decode('utf-8')
    out_bytes = out_bytes.replace("\a","")
    return out_bytes


if __name__ == '__main__':
    import sys 
    #f_to = 'test.docx'
    #html = doc2docx(sys.argv[1], f_to)
    
    text = getTextFromDoc_Libre( sys.argv[1] )
    print("---\n" , text , '\n---')
    #print( html )
    #conv3( sys.argv[1], f_to)
    """
    f = open("a.html", "w")
    f.write( html )
    f.close()
    import office
    """
    



